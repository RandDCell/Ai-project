from docx import Document
from docx.shared import Inches
import os

def create_horizontal_word_document(folder_path, output_file):
    document = Document()
    table = document.add_table(rows=1, cols=1)

    for root, dirs, files in os.walk(folder_path):
        row_cells = table.add_row().cells
        for file in files:
            if file.endswith('.jpg') or file.endswith('.png') or file.endswith('.jpeg'):
                image_path = os.path.join(root, file)
                cell = row_cells[-1]
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(2))
                cell.width = Inches(2)

    document.save(output_file)

# Example usage
folder_path = "C:/Users/randd/OneDrive/Pictures/Screenshots"
output_file = "output1.docx"
create_horizontal_word_document(folder_path, output_file)

